from docx import Document
import re
import os

def split_docx_by_sections(source_filename):
    # 檢查檔案是否存在
    if not os.path.exists(source_filename):
        print(f"找不到檔案: {source_filename}")
        return

    doc = Document(source_filename)
    
    # 儲存分割後的文檔內容
    # 結構: [ {'title': '檔名', 'elements': [段落物件...]} ]
    split_docs = []
    current_elements = []
    current_title = "封面與前言" # 預設第一部分的名稱

    # 定義章節標題的正規表達式 (例如：第一章...第一節...)
    # 根據您的檔案內容，標題格式通常包含 "第一節"、"第二節" 等
    section_pattern = re.compile(r'第[一二三四五六七八九十]+節')

    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 判斷是否為新的一節的標題 (這裡假設標題包含 "第一節" 這樣的字眼且長度適中)
        if section_pattern.search(text) and len(text) < 50:
            # 如果已有累積的內容，先儲存上一個章節
            if current_elements:
                split_docs.append({'title': current_title, 'elements': current_elements})
            
            # 重置，準備下一個章節
            current_title = text.replace('/', '_').replace('\\', '_') # 移除檔名非法字元
            # 簡化檔名，取標題的前20個字
            if len(current_title) > 20:
                current_title = current_title[:20]
            current_elements = [para]
        else:
            current_elements.append(para)

    # 儲存最後一個章節
    if current_elements:
        split_docs.append({'title': current_title, 'elements': current_elements})

    # 將分割後的內容寫入新的檔案
    for item in split_docs:
        new_doc = Document()
        # 這裡簡單地將文字複製過去 (保留格式比較複雜，這裡僅複製文字內容)
        for element in item['elements']:
            new_doc.add_paragraph(element.text)
        
        output_filename = f"{item['title']}.docx"
        new_doc.save(output_filename)
        print(f"已建立檔案: {output_filename}")

# 請將此處改為您檔案的實際路徑
source_file = '健康與護理I 題本.docx' 
split_docx_by_sections(source_file)