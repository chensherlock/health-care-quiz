from docx import Document
import re
import os

def split_docx_by_chapter(source_filename, output_dir='chapters'):
    """
    將 docx 檔案依章節分割成不同檔案
    
    Args:
        source_filename: 來源 docx 檔案路徑
        output_dir: 輸出目錄名稱
    """
    # 檢查檔案是否存在
    if not os.path.exists(source_filename):
        print(f"找不到檔案: {source_filename}")
        return

    # 建立輸出目錄
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document(source_filename)
    
    # 儲存分割後的文檔內容
    # 結構: [ {'title': '檔名', 'elements': [段落物件...]} ]
    split_docs = []
    current_elements = []
    current_title = "封面與前言"  # 預設第一部分的名稱

    # 定義章節標題的正規表達式 (例如：第一章、第二章...)
    chapter_pattern = re.compile(r'第[一二三四五六七八九十]+章')

    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 判斷是否為新的一章的標題
        if chapter_pattern.search(text) and len(text) < 80:
            # 如果已有累積的內容，先儲存上一個章節
            if current_elements:
                split_docs.append({'title': current_title, 'elements': current_elements})
            
            # 重置，準備下一個章節
            # 清理檔名非法字元
            current_title = text.replace('/', '_').replace('\\', '_').replace(':', '_')
            current_title = current_title.replace('*', '_').replace('?', '_').replace('"', '_')
            current_title = current_title.replace('<', '_').replace('>', '_').replace('|', '_')
            current_title = current_title.strip()
            
            # 簡化檔名，取標題的前40個字
            if len(current_title) > 40:
                current_title = current_title[:40]
            current_elements = [para]
        else:
            current_elements.append(para)

    # 儲存最後一個章節
    if current_elements:
        split_docs.append({'title': current_title, 'elements': current_elements})

    # 將分割後的內容寫入新的檔案
    for idx, item in enumerate(split_docs):
        new_doc = Document()
        
        # 複製段落 (包含基本格式)
        for element in item['elements']:
            new_para = new_doc.add_paragraph()
            
            # 複製每個 run 以保留格式
            for run in element.runs:
                new_run = new_para.add_run(run.text)
                # 複製格式
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.name:
                    new_run.font.name = run.font.name
        
        # 產生檔名 (加上序號避免重複)
        output_filename = os.path.join(output_dir, f"{idx:02d}_{item['title']}.docx")
        new_doc.save(output_filename)
        print(f"已建立檔案: {output_filename}")

    print(f"\n完成！共分割成 {len(split_docs)} 個檔案，儲存於 '{output_dir}' 目錄")


if __name__ == '__main__':
    # 來源檔案
    source_file = 'all.docx'
    
    # 執行分割
    split_docx_by_chapter(source_file)
